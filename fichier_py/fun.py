import pandas as pd
from docx.shared import *

# function de modifier types et netwayes chaine de caractaire
# --> type date
def dateType(df, columns):
    for col in columns:
        if col in df.columns:
            df[col] = pd.to_datetime( df[col], format="%d/%m/%Y", errors="coerce" ) 
    return df
# --> type Int
def IntType(df, columns):
    for col in columns:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce").astype("Int64")
    return df
# --> type category
def CategoryType(df, columns):
    for col in columns:
        if col in df.columns:
            df[col] = df[col].astype(str).str.lower().str.strip()
            df[col] = df[col].astype("category")
    return df
# --> type string 
def StringType(df , columns) :
    for col in columns:
        if col in df.columns:
            df[col] = df[col].astype(str).str.lower().str.strip()
    return df

# add section in fichier word
def add_section(doc , title, df):
        doc.add_heading(title, 3)
        p = doc.add_paragraph()
        run = p.add_run(df.head().to_string())
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        info = doc.add_paragraph(f"Nombre de lignes = {len(df)}")
        info.runs[0].italic = True
        doc.add_paragraph("")