from docx import Document
from docx.shared import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import os


def create_report_word(df , df_A , df_Transformation , df_anormalis ,Ventes_Catégorie , 
            Ventes_region , profit_segment , top_product , G_moins , filename="fichier_livrable/rapport.docx"):

    pd.set_option("display.max_columns", None)

    doc = Document()

    title = doc.add_heading("RAPPORT FINAL", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_section(title, df):
        doc.add_heading(title, 3)
        p = doc.add_paragraph()
        run = p.add_run(df.head().to_string())
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        info = doc.add_paragraph(f"Nombre de lignes = {len(df)}")
        info.runs[0].italic = True
        doc.add_paragraph("")

    # columns netwayer et ajouter
    doc.add_heading("types des columns netwayer et columns ajouter : " , 3)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Colorful List'

    header = table.rows[0].cells
    header[0].text = "Types des colonnes — DataFrame brut\n"
    header[1].text = "Types des colonnes — DataFrame nettoyé\n"

    cell1 = table.rows[0].cells[0]
    cell2 = table.rows[0].cells[1]
    #cell1.paragraphs[0].add_run("Types des colonnes — DataFrame brut\n")
    cell1.paragraphs[0].add_run(df.dtypes.to_string())
    #cell2.paragraphs[0].add_run("Types des colonnes — DataFrame nettoyé\n")
    cell2.paragraphs[0].add_run(df_A.dtypes.to_string())


    add_section(
        "Les transformation :" , df_Transformation
    )

    add_section(
        "data anormalis :" , df_anormalis
    )

    add_section(
        "les ventes par catégorie :" , Ventes_Catégorie
    )
    doc.add_picture("fichier_livrable/Ventes_Catégorie.png" , width=Inches(5))
    add_section(
        "les ventes par region :" , Ventes_region
    )
    doc.add_picture("fichier_livrable/Ventes_region.png" , width=Inches(5))
    add_section(
        "les profit par segment :" , profit_segment
    )
    doc.add_picture("fichier_livrable/Profit_segment.png" , width=Inches(5))
    add_section(
        "les top produit :" , top_product
    )
    doc.add_picture("fichier_livrable/Top_produits.png" , width=Inches(5))
    add_section(
        "classement des moins" , G_moins
    )
    doc.add_picture("fichier_livrable/Croissance_mensuelle.png" , width=Inches(5))

        
    doc.save(filename)

    # remove les images
    images = [
        "fichier_livrable/Ventes_Catégorie.png",
        "fichier_livrable/Ventes_region.png",
        "fichier_livrable/Profit_segment.png",
        "fichier_livrable/Top_produits.png",
        "fichier_livrable/Croissance_mensuelle.png"
    ]
    for img in images:
        if os.path.exists(img):
            os.remove(img)