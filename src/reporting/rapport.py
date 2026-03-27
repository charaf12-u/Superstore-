from docx import Document
from src.cleaning.fun import *
from docx.shared import *
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import os


def create_report_word(df , df_A , df_Transformation , df_anormalis ,Ventes_Catégorie , 
            Ventes_region , profit_segment , top_product , G_moins , filename="docs/reports/rapport.docx"):

    pd.set_option("display.max_columns", None)

    doc = Document()

    title = doc.add_heading("RAPPORT FINAL", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER


    # columns netwayer et ajouter
    doc.add_heading("types des columns netwayer et columns ajouter : " , 3)
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Colorful List'

    header = table.rows[0].cells
    header[0].text = "Types des colonnes — DataFrame brut\n"
    header[1].text = "Types des colonnes — DataFrame nettoyé\n"

    cell1 = table.rows[0].cells[0]
    cell2 = table.rows[0].cells[1]
    cell1.paragraphs[0].add_run(df.dtypes.to_string())
    cell2.paragraphs[0].add_run(df_A.dtypes.to_string())


    add_section(
        doc , "Les transformation :" , df_Transformation
    )

    add_section(
        doc , "data anormalis :" , df_anormalis
    )

    add_section(
        doc , "les ventes par catégorie :" , Ventes_Catégorie
    )
    doc.add_picture("data/processed/Ventes_Catégorie.png" , width=Inches(5))
    add_section(
        doc , "les ventes par region :" , Ventes_region
    )
    doc.add_picture("data/processed/Ventes_region.png" , width=Inches(5))
    add_section(
        doc , "les profit par segment :" , profit_segment
    )
    doc.add_picture("data/processed/Profit_segment.png" , width=Inches(5))
    add_section(
        doc , "les top produit :" , top_product
    )
    doc.add_picture("data/processed/Top_produits.png" , width=Inches(5))
    add_section(
        doc , "classement des moins" , G_moins
    )
    doc.add_picture("data/processed/Croissance_mensuelle.png" , width=Inches(5))

        
    doc.save(filename)

    # remove les images
    images = [
        "data/processed/Ventes_Catégorie.png",
        "data/processed/Ventes_region.png",
        "data/processed/Profit_segment.png",
        "data/processed/Top_produits.png",
        "data/processed/Croissance_mensuelle.png"
    ]
    for img in images:
        if os.path.exists(img):
            os.remove(img)